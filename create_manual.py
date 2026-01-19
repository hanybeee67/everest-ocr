
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='Malgun Gothic', size=11):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def set_nepali_font(run, size=12):
    # Windows standard Nepali fonts: Mangal, Nirmala UI
    font_name = 'Mangal' 
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size)

def create_kr_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('에베레스트 멤버십 직원 매뉴얼', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')

    # Section 1
    doc.add_heading('1. 시스템 소개', level=1)
    p = doc.add_paragraph()
    run = p.add_run('이 시스템은 손님이 영수증을 사진으로 찍어 올리면 자동으로 포인트가 적립되는 서비스입니다.\n손님은 적립된 포인트로 메뉴 무료 쿠폰(사모사, 커리 등)을 발급받을 수 있습니다.')
    set_font(run)

    # Section 2
    doc.add_heading('2. 쿠폰 사용 처리 방법', level=1)
    p = doc.add_paragraph()
    run = p.add_run('손님이 쿠폰을 사용하겠다고 스마트폰 화면을 보여줄 때의 절차입니다.')
    set_font(run)

    # Steps
    doc.add_heading('Step 1. 쿠폰 확인', level=2)
    p = doc.add_paragraph()
    run = p.add_run('손님의 폰에서 "사용 가능(AVAILABLE)" 상태인지 확인하세요.')
    set_font(run)

    doc.add_heading('Step 2. 사용 버튼 클릭', level=2)
    p = doc.add_paragraph()
    run = p.add_run('직원이 직접 손님의 폰화면에 있는 [직원에게 보여주기 (사용하기)] 버튼을 누르세요.')
    set_font(run)

    doc.add_heading('Step 3. 지점 선택 및 PIN 입력', level=2)
    p = doc.add_paragraph()
    t = '화면에 팝업창이 뜨면 다음을 수행하세요:\n1. 현재 근무 중인 지점(예: 동대문, 영등포 등)을 선택하세요.\n2. 본인의 직원 PIN 번호(4자리)를 손님의 폰에 입력하고 [승인]을 누르세요.'
    run = p.add_run(t)
    set_font(run)

    doc.add_heading('Step 4. 완료 확인', level=2)
    p = doc.add_paragraph()
    run = p.add_run('화면이 새로고침되면서 "사용완료(USED)"로 바뀌면 처리가 끝난 것입니다. 음식을 서빙주세요.')
    set_font(run)

    # Section 3
    doc.add_heading('3. 중요 주의사항', level=1)
    items = [
        "PIN 번호는 절대 손님에게 알려주지 마세요.",
        "PIN 번호가 기억나지 않으면 매니저에게 문의하세요.",
        "손님이 영수증 적립이 안 된다고 하면, 영수증이 구겨지거나 잘리지 않았는지 확인해주세요."
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_font(run)

    doc.save('Everest_Manual_KR.docx')
    print("Korean Manual created: Everest_Manual_KR.docx")

def create_np_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('एभरेस्ट सदस्यता प्रणाली कर्मचारी म्यानुअल', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set font for title
    for run in title.runs:
        set_nepali_font(run, size=18)
    
    doc.add_paragraph('\n')

    # Section 1
    h1 = doc.add_heading('१. प्रणाली परिचय (System Introduction)', level=1)
    for run in h1.runs: set_nepali_font(run, 14)

    p = doc.add_paragraph()
    text = 'यो प्रणाली एक सेवा हो जहाँ ग्राहकहरूले आफ्नो रसिदको तस्बिर अपलोड गर्दा स्वचालित रूपमा अंक (Points) प्राप्त गर्छन्। ग्राहकहरूले जम्मा गरेको अंक प्रयोग गरेर नि:शुल्क मेनु कुपनहरू (जस्तै समोसा, करी, आदि) प्राप्त गर्न सक्छन्।'
    run = p.add_run(text)
    set_nepali_font(run)

    # Section 2
    h2 = doc.add_heading('२. कुपन कसरी प्रयोग गर्ने (How to Redeem)', level=1)
    for run in h2.runs: set_nepali_font(run, 14)

    p = doc.add_paragraph()
    run = p.add_run('यहाँ प्रक्रिया छ जब ग्राहकले कुपन प्रयोग गर्न आफ्नो स्मार्टफोन स्क्रिन देखाउँछन्।')
    set_nepali_font(run)

    # Steps
    s1 = doc.add_heading('Step 1. कुपन जाँच गर्नुहोस्', level=2)
    for run in s1.runs: set_nepali_font(run, 13)
    
    p = doc.add_paragraph()
    run = p.add_run('ग्राहकको फोनमा कुपन "AVAILABLE" (उपलब्ध) अवस्थामा छ कि छैन जाँच्नुहोस्।')
    set_nepali_font(run)

    s2 = doc.add_heading('Step 2. प्रयोग बटन थिच्नुहोस्', level=2)
    for run in s2.runs: set_nepali_font(run, 13)

    p = doc.add_paragraph()
    run = p.add_run('कर्मचारीले ग्राहकको फोन स्क्रिनमा रहेको [Use This Coupon] बटन थिच्नुहोस्।')
    set_nepali_font(run)

    s3 = doc.add_heading('Step 3. शाखा चयन र PIN कोड', level=2)
    for run in s3.runs: set_nepali_font(run, 13)

    p = doc.add_paragraph()
    text = 'पप-अप विन्डोमा:\n१. हालको शाखा (Branch) चयन गर्नुहोस् (जस्तै: Dongdaemun)।\n२. ग्राहकको फोनमा ४-अंकीय कर्मचारी PIN कोड प्रविष्ट गर्नुहोस् र [Confirm] थिच्नुहोस्।'
    run = p.add_run(text)
    set_nepali_font(run)

    s4 = doc.add_heading('Step 4. पूरा भयो', level=2)
    for run in s4.runs: set_nepali_font(run, 13)

    p = doc.add_paragraph()
    run = p.add_run('स्क्रिन रिफ्रेस भएपछि र "USED" (प्रयोग भयो) देखिपछि प्रक्रिया पूरा हुन्छ। त्यसपछि खाना सर्भ गर्नुहोस्।')
    set_nepali_font(run)

    # Section 3
    h3 = doc.add_heading('३. महत्त्वपूर्ण जानकारी (Important)', level=1)
    for run in h3.runs: set_nepali_font(run, 14)

    items = [
        "तपाइँको PIN कोड ग्राहकहरूलाई कहिल्यै नदिनुहोस्।",
        "यदि तपाइँ आफ्नो PIN कोड बिर्सनुभयो भने, प्रबन्धक (Manager) लाई सोध्नुहोस्।",
        "यदि ग्राहकले रसिद अपलोड काम गरिरहेको छैन भन्छन् भने, रसिद प्रष्ट छ कि छैन जाँच्नुहोस्।"
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_nepali_font(run)

    doc.save('Everest_Manual_NP.docx')
    print("Nepali Manual created: Everest_Manual_NP.docx")

if __name__ == "__main__":
    create_kr_manual()
    create_np_manual()
