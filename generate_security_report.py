
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

def set_font(run, font_name='Malgun Gothic', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

def create_report():
    doc = Document()
    
    # --- Title Page ---
    for i in range(5): doc.add_paragraph() # Spacing
    
    title = doc.add_heading('에베레스트 멤버십(Everest Membership)\n보안 및 안정성 기술 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('\nSecurity Architecture & Data Protection Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i in range(15): doc.add_paragraph() # Spacing

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(info.add_run('작성일: 2026년 1월 28일\n'), size=11)
    set_font(info.add_run('수신: 경영진 및 운영 위원회\n'), size=11)
    set_font(info.add_run('작성자: 기술 개발팀'), size=11)

    doc.add_page_break()

    # --- 1. Executive Summary ---
    h1 = doc.add_heading('1. 개요 (Executive Summary)', level=1)
    p = doc.add_paragraph()
    set_font(p.add_run('본 보고서는 에베레스트 멤버십 시스템에 적용된 보안 조치와 기술적 안전장치를 종합적으로 설명합니다. '), size=11)
    set_font(p.add_run('현재 시스템은 금융권 수준의 암호화, 글로벌 표준 네트워크 보안(Cloudflare), 그리고 안전한 알림 채널(Aligo KakaoTalk)을 통합하여 '), size=11)
    set_font(p.add_run('고객 데이터 보호와 시스템 안정성을 최우선으로 설계되었습니다.'), size=11, bold=True)
    
    # --- 2. Data Protection ---
    h2 = doc.add_heading('2. 데이터 보호 및 암호화 (Data Protection)', level=1)
    
    # 2.1 Encryption
    doc.add_heading('2.1. 민감 정보 암호화 (PII Encryption)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run('적용 대상: '), bold=True)
    set_font(p.add_run('회원 이름, 전화번호, 생년월일 등 모든 개인 식별 정보'))
    
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run('기술 표준: '), bold=True)
    set_font(p.add_run('Fernet 대칭키 암호화 (AES-128 CBC + HMAC-SHA256)'))
    
    p = doc.add_paragraph()
    set_font(p.add_run('  - 데이터베이스가 탈취되더라도, 암호화 키 없이는 고객 정보를 절대 해독할 수 없습니다.\n'), size=10)
    set_font(p.add_run('  - 암호화 키는 소스 코드와 분리되어 안전한 환경 변수로 관리됩니다.'), size=10)

    # 2.2 Blind Indexing
    doc.add_heading('2.2. 안전한 검색 기술 (Secure Blind Indexing)', level=2)
    p = doc.add_paragraph()
    set_font(p.add_run('전화번호로 회원을 검색할 때, 전화번호 원본을 절대 복호화하지 않고 검색하는 최신 보안 기법을 도입했습니다.'))
    
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run('기술 원리: '), bold=True)
    set_font(p.add_run('전화번호 + 고유의 "Pepper" 값 -> SHA-256 단방향 해싱'))
    
    p = doc.add_paragraph()
    set_font(p.add_run('  - 이를 통해 관리자조차도 데이터베이스 내부를 들여다볼 때 고객 전화번호를 알 수 없습니다.'), size=10)

    # 2.3 Password Hashing
    doc.add_heading('2.3. 관리자 및 PIN 보안', level=2)
    p = doc.add_paragraph()
    set_font(p.add_run('비밀번호와 직원 PIN 번호는 저장 즉시 복구가 불가능한 형태로 변환됩니다.'))
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run('기술 표준: '), bold=True)
    set_font(p.add_run('Bcrypt (Salt 포함, Rainbow Table 공격 방어)'))

    # --- 3. Network & Infrastructure ---
    h3 = doc.add_heading('3. 네트워크 및 인프라 보안 (Network Security)', level=1)
    
    # 3.1 Cloudflare
    doc.add_heading('3.1. 전역 보안 네트워크 (Cloudflare)', level=2)
    p = doc.add_paragraph()
    set_font(p.add_run('전 세계 330개 이상의 도시에 위치한 Cloudflare 네트워크를 통해 서비스를 제공합니다.'))
    
    items = [
        ('DDoS 방어', '대량의 트래픽 공격으로부터 서버를 24시간 보호합니다.'),
        ('Web Application Firewall (WAF)', 'SQL Injection, XSS 등 해킹 시도를 자동으로 차단합니다.'),
        ('TLS/SSL 암호화', '고객과 서버 간의 모든 통신은 최신 암호화 프로토콜로 보호됩니다 (Lock Icon).')
    ]
    for key, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        set_font(p.add_run(f'{key}: '), bold=True)
        set_font(p.add_run(desc))

    # 3.2 Rate Limiting
    doc.add_heading('3.2. 비정상 접근 차단 (Rate Limiting)', level=2)
    p = doc.add_paragraph()
    set_font(p.add_run('특정 IP에서 비정상적으로 많은 요청(예: 무작위 비밀번호 대입)이 발생하면 자동으로 접속을 차단하는 시스템이 가동 중입니다.'))

    # --- 4. Notification Security ---
    h4 = doc.add_heading('4. 알림 시스템 보안 (Notification Security)', level=1)
    
    doc.add_heading('4.1. 카카오 알림톡 (Aligo Integration)', level=2)
    p = doc.add_paragraph()
    set_font(p.add_run('고객에게 발송되는 쿠폰 및 안내 메시지는 공식 비즈니스 채널인 "알리고(Aligo)"를 통해 카카오톡으로 전송됩니다.'))
    
    items = [
        ('신뢰성', '스팸 문자가 아닌, 인증된 비즈니스 프로필로 전송되어 고객 신뢰도를 높입니다.'),
        ('보안성', '발송되는 모든 메시지 내용은 전송 구간에서 암호화 처리됩니다.')
    ]
    for key, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        set_font(p.add_run(f'{key}: '), bold=True)
        set_font(p.add_run(desc))

    # --- 5. Conclusion ---
    h5 = doc.add_heading('5. 결론 (Conclusion)', level=1)
    p = doc.add_paragraph()
    set_font(p.add_run('에베레스트 멤버십 시스템은 기획 단계에서부터 "보안(Security by Design)"을 핵심 원칙으로 개발되었습니다. '))
    set_font(p.add_run('현재 적용된 보안 수준은 일반적인 소규모 멤버십 서비스를 훨씬 상회하는 수준이며, '))
    set_font(p.add_run('경영진과 고객 모두가 안심하고 사용할 수 있는 견고한 보안 환경을 갖추고 있음을 보고드립니다.'))
    
    doc.add_paragraph('\n')
    
    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(sig.add_run('(인) 에베레스트 멤버십 개발팀'), size=12, bold=True)

    # Save
    filename = "Everest_Security_Report.docx"
    doc.save(filename)
    print(f"Successfully created {filename}")

if __name__ == "__main__":
    create_report()
