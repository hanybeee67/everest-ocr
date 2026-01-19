
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('에베레스트 멤버십 시스템\n보안 강화 및 안정화 리포트', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Intro
    intro = doc.add_paragraph('고객님의 소중한 개인정보를 보호하고, 더욱 안정적인 서비스를 제공하기 위해 ')
    intro.add_run('첨단 보안 시스템').bold = True
    intro.add_run('을 구축했습니다. 학부모님들께서 안심하고 사용하실 수 있도록 보이지 않는 곳까지 철저하게 점검하고 강화했습니다.')
    
    doc.add_paragraph('\n')

    # Section 1
    h1 = doc.add_heading('1. 개인정보 절대 보호 (Privacy Protection)', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('✅ 영수증 이미지 "즉시 파기" 시스템 도입\n').bold = True
    p1.add_run('업로드해주신 영수증 사진은 OCR(글자 인식) 처리가 끝나는 ')
    p1.add_run('즉시 서버에서 영구 삭제').bold = True
    p1.add_run('됩니다. 처리 후 0.1초도 서버에 남기지 않아 유출 가능성을 원천 차단했습니다.')

    # Section 2
    h2 = doc.add_heading('2. 외부 공격 철통 방어 (Iron Wall Security)', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('✅ 스마트 침입 탐지 및 차단 (Rate Limiting)\n').bold = True
    p2.add_run('누군가 고의로 시스템에 과부하를 주거나 정보를 빼내려는 시도를 하면, 시스템이 이를 ')
    p2.add_run('자동으로 감지하고 즉시 차단').bold = True
    p2.add_run('합니다. (전화번호 조회, 로그인 시도 등 모든 접근을 24시간 감시)')
    
    p2_2 = doc.add_paragraph()
    p2_2.add_run('✅ 관리자 페이지 요새화\n').bold = True
    p2_2.add_run('관리자 접속 경로를 암호화하여 숨기고, 비밀번호는 ')
    p2_2.add_run('군사 등급 수준의 암호화 기술').bold = True
    p2_2.add_run('로 보호되고 있습니다.')

    # Section 3
    h3 = doc.add_heading('3. 빈틈없는 운영 원칙 (System Reliability)', level=1)
    p3 = doc.add_paragraph()
    p3.add_run('✅ 중복 적립 자동 방지 AI\n').bold = True
    p3.add_run('실수로 중복해서 올리거나, 같은 영수증을 다시 사용하는 경우를 시스템이 정확히 걸러냅니다. 모든 회원님께 공정한 혜택이 돌아가도록 설계되었습니다.')

    p3_2 = doc.add_paragraph()
    p3_2.add_run('✅ 365일 무중단 보안 감시\n').bold = True
    p3_2.add_run('운영 환경(Production) 구축을 완료하여, 개발용 코드가 아닌 최적화된 정식 서버 모드로 동작합니다. 불필요한 정보 노출을 막아 시스템이 더욱 빠르고 안전해졌습니다.')

    doc.add_paragraph('\n')
    
    # Footer
    footer = doc.add_paragraph('에베레스트 멤버십 개발팀 드림')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    filename = "Everest_Security_Report.docx"
    doc.save(filename)
    print(f"Successfully created {filename}")

if __name__ == "__main__":
    create_report()
